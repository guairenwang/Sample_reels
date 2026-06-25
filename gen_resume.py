from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

# === Header ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('吴明康')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('嵌入式工程师 / MCU 开发  |  应届生')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('📞 13372103279  |  📧 1656280667@qq.com  |  📍 南京')
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()  # spacer

# === 教育背景 ===
h = doc.add_heading('教育背景', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

p = doc.add_paragraph()
run = p.add_run('南京工程学院')
run.bold = True
run.font.size = Pt(11)
run = p.add_run('    自动化 · 本科    2022.07 - 2026.06')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
run = p.add_run('主修课程：')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('电路原理、模拟电子技术、数字电子技术、自动控制原理、电力电子技术、电机与拖动基础、电气控制与PLC、运动控制系统')
run.font.size = Pt(9.5)

doc.add_paragraph()  # spacer

# === 专业技能 ===
h = doc.add_heading('专业技能', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

skills_text = [
    'MCU开发：精通STM32全系列（F1/F4/C8T6等）开发，熟练使用标准库与HAL库，掌握FreeRTOS实时操作系统',
    '硬件设计：熟练使用嘉立创EDA，具备PCB焊接调试、电路原理图校对、飞线连接与排故能力',
    '编程语言：C / C++、Python、Lua',
    '通信与电机：UART / I2C / SPI、PWM、步进电机驱动、M2006无刷电机控制',
    'PLC与电气：西门子S7-1200系列PLC程序调试，AutoCAD电气图纸绘制，变频器与继电器选型',
    '工具与平台：Git版本控制、Altium Designer、SolidWorks基础、Linux基础'
]
for skill in skills_text:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(skill)
    run.font.size = Pt(10)

doc.add_paragraph()

# === 实习经历 ===
h = doc.add_heading('实习经历', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# 实习1
p = doc.add_paragraph()
run = p.add_run('昆山市禾物数字科技有限公司')
run.bold = True
run.font.size = Pt(11)
run = p.add_run('    技术助理（实习）    2025.07 - 2025.08')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

exp1_details = [
    '毕设文档编写与排版：根据客户提供的项目资料，协助完成毕业设计论文的格式排版、图表优化及参考文献整理，累计服务客户15+',
    'PCB焊接与接线：使用嘉立创EDA查看PCB设计文件，完成元器件手工焊接及飞线连接，配合进行电路通断测试和故障排查',
    '售前售后支持：通过微信/电话对接客户，解答毕设文档修改、电路板使用等基础问题，记录并跟进售后工单',
    '电路原理图校对：协助核对原理图与PCB Layout的一致性，检查封装匹配性及网络连接'
]
for d in exp1_details:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(d)
    run.font.size = Pt(9.5)

doc.add_paragraph()

# 实习2
p = doc.add_paragraph()
run = p.add_run('张家港市创为自动化工程有限公司')
run.bold = True
run.font.size = Pt(11)
run = p.add_run('    自动化助理工程师（实习）    2025.03')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

exp2_details = [
    'PLC程序辅助调试：协助完成西门子S7-1200系列PLC程序下载与在线监控，参与自动配料系统控制逻辑现场调试',
    '电气图纸整理归档：使用AutoCAD整理项目电气原理图及接线图，完成3个非标自动化设备项目图纸标准化归档',
    '现场安装支持：跟随工程师前往客户工厂，参与电气控制柜现场安装与布线，协助设备上电前线路检测',
    '技术文档编写：编写设备操作说明书及维护手册，整理调试日志和故障处理记录'
]
for d in exp2_details:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(d)
    run.font.size = Pt(9.5)

doc.add_paragraph()

# === 项目经历 ===
h = doc.add_heading('项目经历', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# 项目1
p = doc.add_paragraph()
run = p.add_run('双臂魔方机器人')
run.bold = True
run.font.size = Pt(11)
run = p.add_run('    | 毕业设计')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
run = p.add_run('技术栈：')
run.bold = True
run.font.size = Pt(9.5)
run = p.add_run('树莓派Pico、树莓派、TMC2209步进电机驱动、OpenCV')
run.font.size = Pt(9.5)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('下位机采用树莓派Pico，通过TMC2209步进电机驱动控制双臂运动')
run.font.size = Pt(9.5)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('上位机采用树莓派，利用OpenCV识别魔方六个面所有色块，求解还原算法并与下位机通信完成魔方还原')
run.font.size = Pt(9.5)

doc.add_paragraph()

# 项目2
p = doc.add_paragraph()
run = p.add_run('基于Django框架的后台管理系统')
run.bold = True
run.font.size = Pt(11)
run = p.add_run('    | 课程设计')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
run = p.add_run('技术栈：')
run.bold = True
run.font.size = Pt(9.5)
run = p.add_run('Python、Django、JavaScript、HTML、CSS')
run.font.size = Pt(9.5)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('实现了简单的人员增删改查、验证码登录校验、数据可视化功能')
run.font.size = Pt(9.5)

doc.add_paragraph()

# === 竞赛获奖 ===
h = doc.add_heading('竞赛获奖', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

awards = [
    ['2023 中国机器人大赛暨RoboCup机器人世界杯中国赛 全国赛', '二等奖', '负责下位机 — STM32F1 HAL库，实现车辆精准位置控制'],
    ['2023 中国机器人大赛暨RoboCup机器人世界杯中国赛 专项赛', '三等奖', '负责下位机 — STM32C8T6 标准库，实现车辆精准位置控制'],
    ['2023 江苏省大学生工程实践与创新能力大赛', '二等奖', '负责下位机 — STM32F4 HAL库 + WIT陀螺仪，实现车辆位置控制、机械臂控制、物块识别'],
    ['RoboMaster 2024 机甲大师高校联盟赛（山东站）步兵对抗赛', '三等奖', '负责下位机 — C板 HAL库 + FreeRTOS + M2006无刷电机，车辆遥控与炮口射击']
]
for name, level, role in awards:
    p = doc.add_paragraph()
    run = p.add_run(f'[{level}] ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x55, 0x33)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(role)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# === 专利 ===
h = doc.add_heading('专利', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

p = doc.add_paragraph()
run = p.add_run('一种舵轮全地形底盘')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('    | 实用新型专利')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# === 自我评价 ===
h = doc.add_heading('自我评价', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

p = doc.add_paragraph()
run = p.add_run('熟悉STM32各系列开发，能使用标准库与HAL库完成下位机程序编写，接触过FreeRTOS实时操作系统与常见传感器、电机驱动等外设调试。参与过多次国家级、省级机器人竞赛的下位机开发工作，对嵌入式系统的软硬件联调有一定实践经验，具备基本的代码调试与问题排查能力。工作态度踏实，愿意从具体任务做起，仍在持续学习以提升技术水平。')
run.font.size = Pt(10)

# Save
output_path = os.path.expanduser('~/Desktop/AI/ceshi6/简历2.docx')
doc.save(output_path)
print(f'简历已生成: {output_path}')
